"""Narralytica: Export endpoints for SRT, VTT, JSON, CSV, DOCX, and the premium Metadata Script PDF."""
import csv
import io
import json
from datetime import timedelta

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import HexColor, Color
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    KeepTogether, HRFlowable,
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

from app.core.database import get_db
from app.models.models import Episode, Scene, Show, SceneObject

router = APIRouter(prefix="/export", tags=["export"])


def _fmt_ts(seconds: float) -> str:
    total_seconds = int(seconds)
    h, m, s = total_seconds // 3600, (total_seconds % 3600) // 60, total_seconds % 60
    ms = int((seconds - int(seconds)) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _fmt_vtt(seconds: float) -> str:
    total_seconds = int(seconds)
    h, m, s = total_seconds // 3600, (total_seconds % 3600) // 60, total_seconds % 60
    ms = int((seconds - int(seconds)) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"


def _fmt_tc(seconds: float) -> str:
    total_seconds = int(seconds)
    h, m, s = total_seconds // 3600, (total_seconds % 3600) // 60, total_seconds % 60
    return f"{h:02d}:{m:02d}:{s:02d}"


async def _get_episode_scenes(episode_id: int, db: AsyncSession):
    ep = await db.execute(
        select(Episode).where(Episode.id == episode_id).options(selectinload(Episode.show))
    )
    episode = ep.scalar_one_or_none()
    if not episode:
        raise HTTPException(404, "Episode not found")
    result = await db.execute(
        select(Scene).where(Scene.episode_id == episode_id)
        .options(selectinload(Scene.objects)).order_by(Scene.start_timestamp)
    )
    scenes = result.scalars().all()
    if not scenes:
        raise HTTPException(404, "No scenes found for this episode")
    return episode, scenes


def _safe_filename(episode) -> str:
    return f"S{episode.season:02d}E{episode.episode_number:02d}_{episode.title.replace(' ', '_')}"


# ===== SRT =====

@router.get("/{episode_id}/srt")
async def export_srt(episode_id: int, db: AsyncSession = Depends(get_db)):
    episode, scenes = await _get_episode_scenes(episode_id, db)
    lines, counter = [], 1
    for scene in scenes:
        for entry in (scene.merged_transcript or []):
            lines.extend([
                str(counter),
                f"{_fmt_ts(entry.get('start', scene.start_timestamp))} --> {_fmt_ts(entry.get('end', scene.end_timestamp))}",
                f"[{entry.get('speaker', 'Unknown')}] {entry.get('text', '')}",
                "",
            ])
            counter += 1
    if not lines:
        raise HTTPException(404, "No transcript data")
    return StreamingResponse(
        io.BytesIO("\n".join(lines).encode("utf-8")),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{_safe_filename(episode)}.srt"'},
    )


# ===== VTT =====

@router.get("/{episode_id}/vtt")
async def export_vtt(episode_id: int, db: AsyncSession = Depends(get_db)):
    """Export WebVTT subtitle file."""
    episode, scenes = await _get_episode_scenes(episode_id, db)
    lines = ["WEBVTT", ""]
    for scene in scenes:
        for entry in (scene.merged_transcript or []):
            start = _fmt_vtt(entry.get("start", scene.start_timestamp))
            end = _fmt_vtt(entry.get("end", scene.end_timestamp))
            speaker = entry.get("speaker", "Unknown")
            text = entry.get("text", "")
            lines.extend([f"{start} --> {end}", f"<v {speaker}>{text}", ""])
    if len(lines) <= 2:
        raise HTTPException(404, "No transcript data")
    return StreamingResponse(
        io.BytesIO("\n".join(lines).encode("utf-8")),
        media_type="text/vtt; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{_safe_filename(episode)}.vtt"'},
    )


# ===== Script JSON =====

@router.get("/{episode_id}/script-json")
async def export_script_json(episode_id: int, db: AsyncSession = Depends(get_db)):
    episode, scenes = await _get_episode_scenes(episode_id, db)
    script = {
        "episode": {"title": episode.title, "season": episode.season, "episode_number": episode.episode_number},
        "scenes": [{
            "start": s.start_timestamp, "end": s.end_timestamp,
            "characters": s.characters_present, "transcript": s.merged_transcript or [],
        } for s in scenes],
    }
    return StreamingResponse(
        io.BytesIO(json.dumps(script, indent=2).encode("utf-8")),
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{_safe_filename(episode)}_script.json"'},
    )


# ===== Full Metadata JSON (35+ fields) =====

@router.get("/{episode_id}/metadata-json")
async def export_metadata_json(episode_id: int, db: AsyncSession = Depends(get_db)):
    episode, scenes = await _get_episode_scenes(episode_id, db)
    data = {
        "episode": {"title": episode.title, "season": episode.season, "episode_number": episode.episode_number},
        "scenes": [],
    }
    for s in scenes:
        data["scenes"].append({
            "id": s.id, "start": s.start_timestamp, "end": s.end_timestamp, "duration": s.duration,
            "characters_present": s.characters_present, "key_dialog": s.key_dialog,
            "character_interactions": s.character_interactions,
            "character_motivations_feelings": s.character_motivations_feelings,
            "actions": s.actions, "visual_gags": s.visual_gags, "dialog_based_humor": s.dialog_based_humor,
            "location": s.location, "time_of_day": s.time_of_day, "setting_type": s.setting_type,
            "color_palette": s.color_palette, "lighting": s.lighting,
            "camera_shot_type": s.camera_shot_type, "camera_movement": s.camera_movement,
            "scene_composition": s.scene_composition, "visual_style_notes": s.visual_style_notes,
            "music_present": s.music_present, "music_description": s.music_description,
            "sound_effects": s.sound_effects, "ambient_audio": s.ambient_audio,
            "mood_ambience": s.mood_ambience, "scene_pacing": s.scene_pacing,
            "tone": s.tone, "emotional_arc": s.emotional_arc,
            "tropes_memes": s.tropes_memes, "cultural_references": s.cultural_references,
            "recurring_gags": s.recurring_gags, "plot_significance": s.plot_significance,
            "continuity_notes": s.continuity_notes,
            "explicitness_language": s.explicitness_language, "explicitness_violence": s.explicitness_violence,
            "explicitness_sexual": s.explicitness_sexual, "explicitness_substance": s.explicitness_substance,
            "explicitness_thematic": s.explicitness_thematic,
            "scene_transitions": s.scene_transitions, "text_on_screen": s.text_on_screen,
            "overall_confidence": s.overall_confidence, "description_text": s.description_text,
            "objects": [{"name": o.name, "category": o.category, "prominence": o.prominence} for o in s.objects],
        })
    return StreamingResponse(
        io.BytesIO(json.dumps(data, indent=2).encode("utf-8")),
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{_safe_filename(episode)}_metadata.json"'},
    )


# ===== Metadata CSV (35+ fields) =====

@router.get("/{episode_id}/metadata-csv")
async def export_metadata_csv(episode_id: int, db: AsyncSession = Depends(get_db)):
    episode, scenes = await _get_episode_scenes(episode_id, db)
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "Scene", "Start", "End", "Duration", "Characters", "Location", "Setting",
        "Time of Day", "Tone", "Mood", "Pacing", "Plot Significance",
        "Actions", "Visual Gags", "Dialog Humor", "Camera Shot", "Camera Movement",
        "Lighting", "Music", "Sound Effects",
        "Expl: Language", "Expl: Violence", "Expl: Sexual", "Expl: Substance", "Expl: Thematic",
        "Tropes", "Cultural Refs", "Confidence", "Description",
    ])
    for i, s in enumerate(scenes, 1):
        chars = "; ".join(c["name"] for c in (s.characters_present or []) if isinstance(c, dict))
        writer.writerow([
            i, _fmt_tc(s.start_timestamp), _fmt_tc(s.end_timestamp), round(s.duration, 1),
            chars, s.location or "", s.setting_type or "", s.time_of_day or "",
            s.tone or "", s.mood_ambience or "", s.scene_pacing or "", s.plot_significance or "",
            s.actions or "", s.visual_gags or "", s.dialog_based_humor or "",
            s.camera_shot_type or "", s.camera_movement or "", s.lighting or "",
            s.music_description or "", s.sound_effects or "",
            s.explicitness_language, s.explicitness_violence, s.explicitness_sexual,
            s.explicitness_substance, s.explicitness_thematic,
            "; ".join(s.tropes_memes or []), "; ".join(s.cultural_references or []),
            s.overall_confidence, s.description_text or "",
        ])
    return StreamingResponse(
        io.BytesIO(output.getvalue().encode("utf-8-sig")),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{_safe_filename(episode)}_metadata.csv"'},
    )


# ===== Script DOCX =====

@router.get("/{episode_id}/script-docx")
async def export_script_docx(episode_id: int, db: AsyncSession = Depends(get_db)):
    episode, scenes = await _get_episode_scenes(episode_id, db)
    doc = DocxDocument()
    title = doc.add_heading(level=0)
    run = title.add_run(episode.title)
    run.font.size = Pt(24)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Season {episode.season}, Episode {episode.episode_number}\nGenerated by Narralytica")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()

    for i, scene in enumerate(scenes, 1):
        h = doc.add_heading(level=1)
        h.add_run(f"Scene {i}  [{_fmt_tc(scene.start_timestamp)} - {_fmt_tc(scene.end_timestamp)}]")
        chars = [c["name"] for c in (scene.characters_present or []) if isinstance(c, dict)]
        if chars:
            p = doc.add_paragraph()
            p.add_run("Characters: ").bold = True
            p.add_run(", ".join(chars))
        if scene.location:
            p = doc.add_paragraph()
            p.add_run("Location: ").bold = True
            p.add_run(scene.location)
        if scene.actions:
            p = doc.add_paragraph()
            r = p.add_run(scene.actions)
            r.italic = True
        for entry in (scene.merged_transcript or []):
            p = doc.add_paragraph()
            ts = entry.get("start")
            tc = f" [{_fmt_tc(ts)}]" if ts else ""
            p.add_run(f"{entry.get('speaker', 'Unknown').upper()}{tc}").bold = True
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Pt(36)
            p2.add_run(entry.get("text", ""))
        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{_safe_filename(episode)}_script.docx"'},
    )


# ===== PREMIUM: Metadata Script PDF =====

MOOD_COLORS = {
    "comedic": HexColor("#FFD521"),
    "mixed": HexColor("#8B7355"),
    "dramatic": HexColor("#4A90D9"),
    "tense": HexColor("#D94A4A"),
    "melancholy": HexColor("#6B5B95"),
    "absurd": HexColor("#FF6B35"),
    "heartfelt": HexColor("#FF9999"),
}


def _mood_color(tone: str | None) -> Color:
    if not tone:
        return HexColor("#666666")
    return MOOD_COLORS.get(tone.lower().strip(), HexColor("#666666"))


def _expl_bar_text(value: float) -> str:
    filled = round(value * 10)
    return "=" * filled + "-" * (10 - filled) + f" {round(value * 100)}%"


@router.get("/{episode_id}/metadata-pdf")
async def export_metadata_pdf(episode_id: int, db: AsyncSession = Depends(get_db)):
    """Premium Metadata Script PDF — the annotated director's script.

    Color-coded scene headers by mood, character blocks, dialog with
    timestamps, metadata grids, explicitness ratings, and tropes.
    """
    episode, scenes = await _get_episode_scenes(episode_id, db)
    show = episode.show

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        topMargin=0.6 * inch, bottomMargin=0.5 * inch,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("SceneTitle", parent=styles["Heading2"],
        textColor=HexColor("#FFFFFF"), fontSize=13, spaceAfter=2, spaceBefore=0))
    styles.add(ParagraphStyle("SceneSub", parent=styles["Normal"],
        textColor=HexColor("#EEEEEE"), fontSize=8, spaceAfter=0))
    styles.add(ParagraphStyle("CharBlock", parent=styles["Normal"],
        fontSize=9, textColor=HexColor("#333333"), leading=12))
    styles.add(ParagraphStyle("Dialog", parent=styles["Normal"],
        fontSize=9, leftIndent=24, leading=13, textColor=HexColor("#222222")))
    styles.add(ParagraphStyle("Speaker", parent=styles["Normal"],
        fontSize=9, leading=13, textColor=HexColor("#000000")))
    styles.add(ParagraphStyle("ActionText", parent=styles["Normal"],
        fontSize=9, italic=True, textColor=HexColor("#555555"), leading=12))
    styles.add(ParagraphStyle("MetaLabel", parent=styles["Normal"],
        fontSize=7, textColor=HexColor("#888888"), leading=10))
    styles.add(ParagraphStyle("MetaVal", parent=styles["Normal"],
        fontSize=8, textColor=HexColor("#333333"), leading=11))
    styles.add(ParagraphStyle("SmallCenter", parent=styles["Normal"],
        fontSize=7, textColor=HexColor("#AAAAAA"), alignment=TA_CENTER))

    story = []

    # --- Title page ---
    story.append(Spacer(1, 2 * inch))
    story.append(Paragraph(f"<b>{episode.title}</b>", styles["Title"]))
    story.append(Spacer(1, 12))
    show_name = show.name if show else "Unknown"
    story.append(Paragraph(
        f"Season {episode.season}, Episode {episode.episode_number}<br/>{show_name}",
        ParagraphStyle("TitleSub", parent=styles["Normal"], fontSize=14, alignment=TA_CENTER, textColor=HexColor("#666666"))
    ))
    story.append(Spacer(1, 20))
    story.append(Paragraph(
        f"{len(scenes)} scenes analyzed | Generated by Narralytica",
        ParagraphStyle("TitleMeta", parent=styles["Normal"], fontSize=10, alignment=TA_CENTER, textColor=HexColor("#999999"))
    ))
    story.append(Spacer(1, 3 * inch))

    # --- Each scene ---
    for i, scene in enumerate(scenes, 1):
        elems = []

        # Header bar
        bar_color = _mood_color(scene.tone)
        tc = f"{_fmt_tc(scene.start_timestamp)} - {_fmt_tc(scene.end_timestamp)}"
        loc = scene.location or "Unknown Location"
        header_data = [[
            Paragraph(f"<b>Scene {i}</b>", styles["SceneTitle"]),
            Paragraph(f"{tc}  |  {scene.duration:.0f}s  |  {loc}  |  Conf: {scene.overall_confidence * 100:.0f}%", styles["SceneSub"]),
        ]]
        ht = Table(header_data, colWidths=[1.1 * inch, 5.7 * inch])
        ht.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), bar_color),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ]))
        elems.append(ht)
        elems.append(Spacer(1, 6))

        # Characters
        chars = [c for c in (scene.characters_present or []) if isinstance(c, dict)]
        if chars:
            elems.append(Paragraph(f"<b>Characters:</b> {', '.join(c['name'] for c in chars)}", styles["CharBlock"]))
            elems.append(Spacer(1, 4))

        # Actions
        if scene.actions:
            elems.append(Paragraph(scene.actions, styles["ActionText"]))
            elems.append(Spacer(1, 4))

        # Dialog
        dialog_source = scene.merged_transcript or scene.key_dialog or []
        for entry in dialog_source:
            speaker = entry.get("speaker", "Unknown")
            text = entry.get("text") or entry.get("exact_quote", "")
            ts = entry.get("start") or entry.get("timestamp")
            tc_str = f" [{_fmt_tc(ts)}]" if ts else ""
            elems.append(Paragraph(f"<b>{speaker.upper()}</b>{tc_str}", styles["Speaker"]))
            elems.append(Paragraph(text, styles["Dialog"]))
            elems.append(Spacer(1, 2))

        elems.append(Spacer(1, 6))

        # Metadata grid
        meta = []
        for label, val in [
            ("Tone", scene.tone), ("Mood", scene.mood_ambience), ("Pacing", scene.scene_pacing),
            ("Plot", scene.plot_significance), ("Camera", scene.camera_shot_type),
            ("Movement", scene.camera_movement), ("Lighting", scene.lighting),
            ("Music", scene.music_description), ("Emotional Arc", scene.emotional_arc),
            ("SFX", scene.sound_effects), ("Humor", scene.visual_gags),
        ]:
            if val:
                meta.append([Paragraph(f"<b>{label}</b>", styles["MetaLabel"]), Paragraph(str(val), styles["MetaVal"])])

        if meta:
            mt = Table(meta, colWidths=[1 * inch, 5.8 * inch])
            mt.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), HexColor("#F8F8F8")),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("LINEBELOW", (0, 0), (-1, -2), 0.5, HexColor("#EEEEEE")),
            ]))
            elems.append(mt)
            elems.append(Spacer(1, 4))

        # Explicitness
        has_expl = any([scene.explicitness_language, scene.explicitness_violence,
                        scene.explicitness_sexual, scene.explicitness_substance, scene.explicitness_thematic])
        if has_expl:
            expl_parts = []
            for lbl, val in [("Lang", scene.explicitness_language), ("Violence", scene.explicitness_violence),
                             ("Sexual", scene.explicitness_sexual), ("Substance", scene.explicitness_substance),
                             ("Thematic", scene.explicitness_thematic)]:
                if val > 0:
                    expl_parts.append(f"{lbl}: {_expl_bar_text(val)}")
            elems.append(Paragraph("  |  ".join(expl_parts), styles["MetaLabel"]))
            elems.append(Spacer(1, 4))

        # Tropes / refs
        if scene.tropes_memes:
            elems.append(Paragraph(f"<b>Tropes:</b> {'; '.join(scene.tropes_memes)}", styles["MetaVal"]))
        if scene.cultural_references:
            elems.append(Paragraph(f"<b>Cultural:</b> {'; '.join(scene.cultural_references)}", styles["MetaVal"]))

        elems.append(Spacer(1, 8))
        elems.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#DDDDDD")))
        elems.append(Spacer(1, 12))

        story.append(KeepTogether(elems))

    story.append(Spacer(1, 20))
    story.append(Paragraph("Generated by Narralytica | Natural-Language Video Intelligence", styles["SmallCenter"]))

    doc.build(story)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{_safe_filename(episode)}_metadata_script.pdf"'},
    )
