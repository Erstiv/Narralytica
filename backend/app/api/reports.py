"""Narralytica: Reports — flexible document generator from indexed scene data.

Generates customizable reports at episode or season scope:
  - Plot summaries
  - Character summaries & motivations
  - Two-column scripts (dialog/audio LEFT, visuals/action RIGHT)
  - Scene breakdowns
  - Dialog-only transcripts
  - Visual descriptions

Supports in-browser JSON preview and DOCX download.
"""
import io
import json
import logging
import os
from collections import defaultdict

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.database import get_db
from app.models.models import Episode, Scene, Show

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/reports", tags=["reports"])


# ─── Request / Response Models ──────────────────────────────

VALID_REPORT_TYPES = {
    "plot_summary", "character_summaries", "character_motivations",
    "two_column_script", "scene_breakdown", "dialog_only", "visual_descriptions",
    "content_critic", "marketing_editor",
}


class ReportRequest(BaseModel):
    scope: str  # "episode" or "season"
    episode_id: int | None = None
    show_id: int | None = None
    season: int | None = None
    report_types: list[str]
    nl_query: str | None = None
    format: str = "preview"  # "preview" or "docx"


# ─── Data Fetching ──────────────────────────────────────────

async def _get_episode_data(episode_id: int, db: AsyncSession):
    """Fetch a single episode with its scenes and show."""
    result = await db.execute(
        select(Episode)
        .options(selectinload(Episode.scenes))
        .where(Episode.id == episode_id)
    )
    episode = result.scalar_one_or_none()
    if not episode:
        raise HTTPException(404, f"Episode {episode_id} not found")

    show_result = await db.execute(select(Show).where(Show.id == episode.show_id))
    show = show_result.scalar_one()

    scenes = sorted(episode.scenes, key=lambda s: s.start_timestamp)
    return show, {episode.id: (episode, scenes)}


async def _get_season_data(show_id: int, season: int, db: AsyncSession):
    """Fetch all episodes for a show+season with their scenes."""
    show_result = await db.execute(select(Show).where(Show.id == show_id))
    show = show_result.scalar_one_or_none()
    if not show:
        raise HTTPException(404, f"Show {show_id} not found")

    result = await db.execute(
        select(Episode)
        .options(selectinload(Episode.scenes))
        .where(Episode.show_id == show_id, Episode.season == season)
        .order_by(Episode.episode_number)
    )
    episodes = result.scalars().all()
    if not episodes:
        raise HTTPException(404, f"No episodes for show {show_id} season {season}")

    data = {}
    for ep in episodes:
        scenes = sorted(ep.scenes, key=lambda s: s.start_timestamp)
        data[ep.id] = (ep, scenes)

    return show, data


# ─── Helpers ────────────────────────────────────────────────

def _ep_label(ep: Episode) -> str:
    return f"S{ep.season:02d}E{ep.episode_number:02d}"


def _fmt_time(seconds: float) -> str:
    m = int(seconds // 60)
    s = int(seconds % 60)
    return f"{m}:{s:02d}"


def _safe_filename(show_name: str, label: str, suffix: str) -> str:
    safe = "".join(c if c.isalnum() or c in " _-" else "" for c in f"{show_name}_{label}")
    return f"{safe.strip()}.{suffix}"


def _get_characters_from_scene(scene) -> list[str]:
    return [c.get("name", "Unknown") for c in (scene.characters_present or [])]


# ─── Builder Functions ──────────────────────────────────────

async def build_plot_summary(episodes_scenes: dict) -> dict:
    """AI-generated narrative plot summary — plain English, 2-3 paragraphs per episode,
    plus 1-2 sentence scene summaries."""
    api_key = os.environ.get("GEMINI_API_KEY")

    episodes = []
    for ep_id, (ep, scenes) in episodes_scenes.items():
        # Collect scene descriptions for Gemini
        scene_summaries = []
        for i, scene in enumerate(scenes, 1):
            desc = scene.description_text or scene.actions or ""
            chars = ", ".join(_get_characters_from_scene(scene))
            scene_summaries.append(
                f"Scene {i} [{_fmt_time(scene.start_timestamp)}] ({scene.tone or '?'}, {scene.plot_significance or '?'}): "
                f"Characters: {chars}. {desc[:200]}"
            )

        # Ask Gemini for a proper narrative summary
        narrative = ""
        scene_oneliners = []
        if api_key and scene_summaries:
            try:
                from google import genai
                from google.genai import types
                client = genai.Client(api_key=api_key)
                prompt = (
                    f"You are summarizing Episode '{ep.title}' of a TV show.\n\n"
                    f"Here are all {len(scene_summaries)} scenes:\n" +
                    "\n".join(scene_summaries[:60]) +
                    "\n\nReturn a JSON object with:\n"
                    '1. "narrative": A 2-3 paragraph plain English plot summary telling what happens in this episode as a cohesive story. '
                    'Write it as you would for a TV guide or review — engaging, clear, no jargon.\n'
                    '2. "scene_summaries": An array of 1-2 sentence summaries for each scene, in order.\n\n'
                    'Return ONLY valid JSON: {"narrative": "...", "scene_summaries": ["...", "..."]}'
                )
                response = client.models.generate_content(
                    model="gemini-2.5-flash", contents=prompt,
                    config=types.GenerateContentConfig(temperature=0.3, response_mime_type="application/json"),
                )
                result = json.loads(response.text)
                narrative = result.get("narrative", "")
                scene_oneliners = result.get("scene_summaries", [])
            except Exception as e:
                logger.warning(f"Plot summary Gemini call failed: {e}")

        # Fallback if Gemini fails
        if not narrative:
            narrative = " ".join(
                s.description_text for s in scenes
                if s.description_text and s.plot_significance in ("high", "critical")
            )[:1000]

        episodes.append({
            "episode_label": _ep_label(ep),
            "episode_title": ep.title,
            "narrative": narrative,
            "scene_summaries": scene_oneliners,
        })

    return {"type": "plot_summary", "title": "Plot Summary", "episodes": episodes}


async def build_character_summaries(episodes_scenes: dict) -> dict:
    """AI-generated character profiles — what we know about each character."""
    api_key = os.environ.get("GEMINI_API_KEY")

    # Aggregate raw data per character
    chars = defaultdict(lambda: {
        "scene_count": 0, "total_duration": 0.0,
        "locations": set(), "motivations": [], "speaking_lines": 0,
        "episodes_seen": set(), "tones_seen": defaultdict(int),
        "key_quotes": [], "interactions": [],
    })

    for ep_id, (ep, scenes) in episodes_scenes.items():
        for scene in scenes:
            for c in (scene.characters_present or []):
                name = c.get("name", "Unknown")
                chars[name]["scene_count"] += 1
                chars[name]["total_duration"] += scene.duration
                chars[name]["episodes_seen"].add(_ep_label(ep))
                if scene.location:
                    chars[name]["locations"].add(scene.location)
                if scene.tone:
                    chars[name]["tones_seen"][scene.tone] += 1

            if scene.character_motivations_feelings:
                for c in (scene.characters_present or []):
                    chars[c.get("name", "Unknown")]["motivations"].append(
                        scene.character_motivations_feelings
                    )

            for d in (scene.key_dialog or []):
                speaker = d.get("speaker", "")
                if speaker in chars:
                    chars[speaker]["key_quotes"].append(d.get("quote", "")[:100])

            for t in (scene.merged_transcript or []):
                speaker = t.get("speaker", "Unknown")
                if speaker in chars:
                    chars[speaker]["speaking_lines"] += 1

            for inter in (scene.character_interactions or []):
                for who in [inter.get("character_a"), inter.get("character_b")]:
                    if who and who in chars:
                        chars[who]["interactions"].append(inter.get("description", ""))

    # Build profiles — use Gemini for narrative description
    profiles = []
    significant_chars = [(name, data) for name, data in chars.items() if data["scene_count"] >= 2]
    significant_chars.sort(key=lambda x: -x[1]["total_duration"])

    # Ask Gemini for character descriptions
    char_descriptions = {}
    if api_key and significant_chars:
        try:
            from google import genai
            from google.genai import types
            client = genai.Client(api_key=api_key)

            char_data = "\n".join(
                f"- {name}: {data['scene_count']} scenes, {int(data['total_duration']//60)}m, "
                f"episodes: {', '.join(sorted(data['episodes_seen']))}, "
                f"motivations: {'; '.join(list(dict.fromkeys(data['motivations']))[:3])}, "
                f"quotes: {'; '.join(data['key_quotes'][:3])}, "
                f"interactions: {'; '.join(list(dict.fromkeys(data['interactions']))[:3])}"
                for name, data in significant_chars[:20]
            )

            prompt = (
                f"Based on this character data from a TV show, write a 2-3 sentence profile for each character. "
                f"Describe who they are, what role they play, and their personality/arc as observed.\n\n"
                f"{char_data}\n\n"
                f"Return JSON: {{\"characters\": {{\"Character Name\": \"2-3 sentence description\", ...}}}}"
            )
            response = client.models.generate_content(
                model="gemini-2.5-flash", contents=prompt,
                config=types.GenerateContentConfig(temperature=0.3, response_mime_type="application/json"),
            )
            result = json.loads(response.text)
            char_descriptions = result.get("characters", {})
        except Exception as e:
            logger.warning(f"Character summary Gemini call failed: {e}")

    for name, data in significant_chars[:25]:
        top_tone = max(data["tones_seen"], key=data["tones_seen"].get) if data["tones_seen"] else None
        profiles.append({
            "name": name,
            "description": char_descriptions.get(name, ""),
            "scene_count": data["scene_count"],
            "total_duration": round(data["total_duration"], 1),
            "duration_formatted": f"{int(data['total_duration'] // 60)}m {int(data['total_duration'] % 60)}s",
            "locations": sorted(data["locations"])[:10],
            "episodes": sorted(data["episodes_seen"]),
            "speaking_lines": data["speaking_lines"],
            "dominant_tone": top_tone,
            "key_quotes": data["key_quotes"][:3],
        })

    return {"type": "character_summaries", "title": "Character Summaries", "characters": profiles}


async def build_character_motivations(episodes_scenes: dict) -> dict:
    """Per-character motivations using Gemini to properly parse the combined text."""
    api_key = os.environ.get("GEMINI_API_KEY")

    # Collect all motivation data
    raw_entries = []
    for ep_id, (ep, scenes) in episodes_scenes.items():
        for i, scene in enumerate(scenes, 1):
            motivation_text = scene.character_motivations_feelings or ""
            if not motivation_text:
                continue
            scene_chars = _get_characters_from_scene(scene)
            raw_entries.append({
                "episode": _ep_label(ep),
                "scene": i,
                "time": _fmt_time(scene.start_timestamp),
                "location": scene.location or "Unknown",
                "tone": scene.tone or "?",
                "characters": scene_chars,
                "raw_text": motivation_text,
            })

    if not raw_entries:
        return {"type": "character_motivations", "title": "Character Motivations", "characters": {}}

    # Ask Gemini to split motivations per character
    chars = defaultdict(list)
    if api_key:
        try:
            from google import genai
            from google.genai import types
            client = genai.Client(api_key=api_key)

            # Send in batches of 30 scenes
            for batch_start in range(0, len(raw_entries), 30):
                batch = raw_entries[batch_start:batch_start + 30]
                scenes_text = "\n".join(
                    f"Scene {e['scene']} [{e['time']}] Characters: {', '.join(e['characters'])}. Motivations: {e['raw_text']}"
                    for e in batch
                )
                prompt = (
                    "Split these scene motivation descriptions into individual character motivations.\n"
                    "For each scene, extract what EACH character individually wants or feels.\n\n"
                    f"{scenes_text}\n\n"
                    "Return JSON: {\"scenes\": [{\"scene\": 1, \"characters\": {\"CharName\": \"What they want/feel in 1 sentence\"}}]}"
                )
                response = client.models.generate_content(
                    model="gemini-2.5-flash", contents=prompt,
                    config=types.GenerateContentConfig(temperature=0.2, response_mime_type="application/json"),
                )
                result = json.loads(response.text)

                for parsed_scene in result.get("scenes", []):
                    scene_num = parsed_scene.get("scene", 0)
                    # Find matching raw entry
                    matching = [e for e in batch if e["scene"] == scene_num]
                    if not matching:
                        continue
                    entry = matching[0]

                    for char_name, motivation in parsed_scene.get("characters", {}).items():
                        if motivation and len(motivation) > 5:
                            chars[char_name].append({
                                "episode": entry["episode"],
                                "scene": entry["scene"],
                                "time": entry["time"],
                                "location": entry["location"],
                                "tone": entry["tone"],
                                "motivation": motivation,
                            })

        except Exception as e:
            logger.warning(f"Motivation Gemini parse failed: {e}")

    # Fallback: simple sentence splitting if Gemini failed
    if not chars:
        for entry in raw_entries:
            for char_name in entry["characters"]:
                first = char_name.split()[0].lower()
                sentences = [s.strip() for s in entry["raw_text"].replace(". ", ".|").split("|") if s.strip()]
                char_sents = [s for s in sentences if first in s.lower()]
                if char_sents:
                    chars[char_name].append({
                        "episode": entry["episode"],
                        "scene": entry["scene"],
                        "time": entry["time"],
                        "location": entry["location"],
                        "tone": entry["tone"],
                        "motivation": " ".join(char_sents),
                    })

    sorted_chars = {
        name: entries
        for name, entries in sorted(chars.items(), key=lambda x: -len(x[1]))
        if len(entries) >= 2
    }

    return {
        "type": "character_motivations",
        "title": "Character Motivations",
        "characters": sorted_chars,
    }


def build_two_column_script(episodes_scenes: dict) -> dict:
    """Two-column script: LEFT=dialog/audio, RIGHT=visuals/action."""
    episodes = []
    for ep_id, (ep, scenes) in episodes_scenes.items():
        rows = []
        for i, scene in enumerate(scenes, 1):
            # LEFT: Dialog + Audio
            # Use merged_transcript (complete dialog) as primary source
            # Enhance speaker names from key_dialog where possible
            left_parts = []

            # Build speaker name lookup from key_dialog (Gemini-identified speakers)
            speaker_map = {}
            for d in (scene.key_dialog or []):
                quote_snippet = d.get("quote", "")[:30].lower()
                if d.get("speaker") and quote_snippet:
                    speaker_map[quote_snippet] = d["speaker"]

            chars = _get_characters_from_scene(scene)
            transcript = scene.merged_transcript or []

            for t in transcript:
                speaker = t.get("speaker", "")
                text = t.get("text", "")
                if not text.strip():
                    continue

                # Try to resolve "Unknown" speaker
                if speaker in ("Unknown", ""):
                    # Check key_dialog for a match
                    text_snippet = text[:30].lower()
                    matched = speaker_map.get(text_snippet, "")
                    if matched:
                        speaker = matched
                    elif len(chars) == 1:
                        speaker = chars[0]
                    else:
                        speaker = "—"

                left_parts.append({
                    "speaker": speaker,
                    "text": text,
                    "time": _fmt_time(t.get("start", 0)),
                })

            # If no transcript, fall back to key_dialog only
            if not left_parts:
                for d in (scene.key_dialog or []):
                    if d.get("quote"):
                        left_parts.append({
                            "speaker": d.get("speaker", "—"),
                            "text": d["quote"],
                            "time": _fmt_time(d.get("timestamp", 0)),
                        })

            audio = []
            if scene.music_description:
                audio.append(f"MUSIC: {scene.music_description}")
            if scene.sound_effects:
                audio.append(f"SFX: {scene.sound_effects}")
            if scene.ambient_audio:
                audio.append(f"AMBIENT: {scene.ambient_audio}")

            # RIGHT: Visuals + Action
            right_parts = {
                "description": scene.description_text,
                "actions": scene.actions,
                "camera": f"{scene.camera_shot_type or ''} / {scene.camera_movement or ''}".strip(" /"),
                "composition": scene.scene_composition,
                "lighting": scene.lighting,
                "color_palette": scene.color_palette or [],
                "text_on_screen": scene.text_on_screen,
            }

            rows.append({
                "scene_number": i,
                "time": f"{_fmt_time(scene.start_timestamp)} - {_fmt_time(scene.end_timestamp)}",
                "location": scene.location or "",
                "tone": scene.tone or "",
                "left": {"dialog": left_parts, "audio": audio},
                "right": right_parts,
            })

        episodes.append({
            "episode_label": _ep_label(ep),
            "episode_title": ep.title,
            "rows": rows,
        })

    return {"type": "two_column_script", "title": "Two-Column Script", "episodes": episodes}


def build_scene_breakdown(episodes_scenes: dict) -> dict:
    """Full metadata per scene."""
    episodes = []
    for ep_id, (ep, scenes) in episodes_scenes.items():
        scene_list = []
        for i, scene in enumerate(scenes, 1):
            scene_list.append({
                "scene_number": i,
                "time": f"{_fmt_time(scene.start_timestamp)} - {_fmt_time(scene.end_timestamp)}",
                "duration": f"{scene.duration:.0f}s",
                "characters": _get_characters_from_scene(scene),
                "location": scene.location,
                "tone": scene.tone,
                "mood": scene.mood_ambience,
                "pacing": scene.scene_pacing,
                "emotional_arc": scene.emotional_arc,
                "actions": scene.actions,
                "description": scene.description_text,
                "camera": f"{scene.camera_shot_type or ''} / {scene.camera_movement or ''}".strip(" /"),
                "lighting": scene.lighting,
                "music": scene.music_description,
                "sound_effects": scene.sound_effects,
                "plot_significance": scene.plot_significance,
                "explicitness": {
                    "language": scene.explicitness_language,
                    "violence": scene.explicitness_violence,
                    "sexual": scene.explicitness_sexual,
                    "substance": scene.explicitness_substance,
                    "thematic": scene.explicitness_thematic,
                },
                "tropes": scene.tropes_memes or [],
                "cultural_references": scene.cultural_references or [],
                "continuity": scene.continuity_notes,
                "motivations": scene.character_motivations_feelings,
            })
        episodes.append({
            "episode_label": _ep_label(ep),
            "episode_title": ep.title,
            "scenes": scene_list,
        })

    return {"type": "scene_breakdown", "title": "Scene-by-Scene Breakdown", "episodes": episodes}


def build_dialog_only(episodes_scenes: dict) -> dict:
    """Dialog transcript formatted as speaker:text."""
    episodes = []
    for ep_id, (ep, scenes) in episodes_scenes.items():
        lines = []
        for scene in scenes:
            scene_lines = []
            for t in (scene.merged_transcript or []):
                scene_lines.append({
                    "speaker": t.get("speaker", "Unknown"),
                    "text": t.get("text", ""),
                    "time": _fmt_time(t.get("start", 0)),
                })
            if not scene_lines:
                for d in (scene.key_dialog or []):
                    scene_lines.append({
                        "speaker": d.get("speaker", "Unknown"),
                        "text": d.get("quote", ""),
                        "time": _fmt_time(d.get("timestamp", 0)),
                    })
            if scene_lines:
                lines.append({
                    "scene_time": _fmt_time(scene.start_timestamp),
                    "location": scene.location or "",
                    "dialog": scene_lines,
                })
        episodes.append({
            "episode_label": _ep_label(ep),
            "episode_title": ep.title,
            "scenes": lines,
        })

    return {"type": "dialog_only", "title": "Dialog Transcript", "episodes": episodes}


def build_visual_descriptions(episodes_scenes: dict) -> dict:
    """Visual descriptions: what's on screen."""
    episodes = []
    for ep_id, (ep, scenes) in episodes_scenes.items():
        descs = []
        for i, scene in enumerate(scenes, 1):
            descs.append({
                "scene_number": i,
                "time": f"{_fmt_time(scene.start_timestamp)} - {_fmt_time(scene.end_timestamp)}",
                "location": scene.location,
                "description": scene.description_text,
                "actions": scene.actions,
                "camera": f"{scene.camera_shot_type or ''} / {scene.camera_movement or ''}".strip(" /"),
                "composition": scene.scene_composition,
                "lighting": scene.lighting,
                "visual_style": scene.visual_style_notes,
                "color_palette": scene.color_palette or [],
                "text_on_screen": scene.text_on_screen,
            })
        episodes.append({
            "episode_label": _ep_label(ep),
            "episode_title": ep.title,
            "scenes": descs,
        })

    return {"type": "visual_descriptions", "title": "Visual Descriptions", "episodes": episodes}


def build_content_critic(episodes_scenes: dict) -> dict:
    """Fair but harsh critical analysis derived from indexed scene data.

    Examines: pacing consistency, tonal balance, character equity,
    dialog density, visual variety, plot momentum, and structural issues.
    All opinions grounded in measurable data from the scenes.
    """
    all_scenes = []
    ep_data = []
    char_counts = defaultdict(int)
    char_durations = defaultdict(float)
    char_speaking = defaultdict(int)
    tone_counts = defaultdict(int)
    pacing_counts = defaultdict(int)
    location_counts = defaultdict(int)
    camera_counts = defaultdict(int)
    total_duration = 0
    total_dialog_lines = 0
    scenes_without_dialog = 0
    scenes_without_description = 0
    plot_sig_counts = defaultdict(int)

    for ep_id, (ep, scenes) in episodes_scenes.items():
        ep_tones = defaultdict(int)
        ep_pacings = []
        ep_chars = set()
        ep_dialog_lines = 0
        ep_durations = [s.duration for s in scenes]

        for scene in scenes:
            all_scenes.append(scene)
            total_duration += scene.duration

            # Characters
            for c in (scene.characters_present or []):
                name = c.get("name", "Unknown")
                char_counts[name] += 1
                char_durations[name] += scene.duration
                ep_chars.add(name)

            # Dialog
            transcript = scene.merged_transcript or []
            dialog_count = len(transcript) + len(scene.key_dialog or [])
            total_dialog_lines += dialog_count
            ep_dialog_lines += dialog_count
            if dialog_count == 0:
                scenes_without_dialog += 1

            for t in transcript:
                speaker = t.get("speaker", "Unknown")
                if speaker != "Unknown":
                    char_speaking[speaker] += 1

            # Tone/Pacing
            if scene.tone:
                tone_counts[scene.tone] += 1
                ep_tones[scene.tone] += 1
            if scene.scene_pacing:
                pacing_counts[scene.scene_pacing] += 1
                ep_pacings.append(scene.scene_pacing)

            # Location variety
            if scene.location:
                location_counts[scene.location] += 1
            if scene.camera_shot_type:
                camera_counts[scene.camera_shot_type] += 1

            # Plot significance
            if scene.plot_significance:
                plot_sig_counts[scene.plot_significance] += 1

            if not scene.description_text:
                scenes_without_description += 1

        # Per-episode analysis
        duration_variance = 0
        if ep_durations:
            avg_dur = sum(ep_durations) / len(ep_durations)
            duration_variance = sum((d - avg_dur) ** 2 for d in ep_durations) / len(ep_durations)

        ep_data.append({
            "label": _ep_label(ep),
            "title": ep.title,
            "scene_count": len(scenes),
            "unique_chars": len(ep_chars),
            "dialog_lines": ep_dialog_lines,
            "dominant_tone": max(ep_tones, key=ep_tones.get) if ep_tones else "N/A",
            "tone_variety": len(ep_tones),
            "pacing_shifts": sum(1 for i in range(1, len(ep_pacings)) if ep_pacings[i] != ep_pacings[i-1]),
            "duration_variance": round(duration_variance, 1),
        })

    num_scenes = len(all_scenes)
    num_episodes = len(ep_data)

    # ── Generate critiques ──

    critiques = []

    # 1. Pacing analysis
    if pacing_counts:
        dominant_pacing = max(pacing_counts, key=pacing_counts.get)
        pacing_pct = pacing_counts[dominant_pacing] / num_scenes * 100 if num_scenes else 0
        if pacing_pct > 60:
            critiques.append({
                "category": "Pacing",
                "severity": "warning",
                "finding": f"Monotonous pacing: {pacing_pct:.0f}% of scenes are '{dominant_pacing}'. "
                           f"Effective storytelling typically varies rhythm. Only {len(pacing_counts)} distinct pacing levels used.",
                "data": dict(pacing_counts),
            })
        else:
            critiques.append({
                "category": "Pacing",
                "severity": "strength",
                "finding": f"Good pacing variety with {len(pacing_counts)} distinct levels. "
                           f"No single pacing dominates excessively ({dominant_pacing} at {pacing_pct:.0f}%).",
                "data": dict(pacing_counts),
            })

    # 2. Tonal balance
    if tone_counts:
        dominant_tone = max(tone_counts, key=tone_counts.get)
        tone_pct = tone_counts[dominant_tone] / num_scenes * 100 if num_scenes else 0
        if tone_pct > 50 and len(tone_counts) < 4:
            critiques.append({
                "category": "Tonal Range",
                "severity": "warning",
                "finding": f"Limited tonal range: '{dominant_tone}' dominates at {tone_pct:.0f}% of scenes. "
                           f"Only {len(tone_counts)} distinct tones detected. This may feel one-note to audiences.",
                "data": dict(tone_counts),
            })
        else:
            critiques.append({
                "category": "Tonal Range",
                "severity": "strength",
                "finding": f"Healthy tonal variety with {len(tone_counts)} distinct tones. "
                           f"'{dominant_tone}' leads at {tone_pct:.0f}% but doesn't overwhelm.",
                "data": dict(tone_counts),
            })

    # 3. Character equity
    if char_counts:
        sorted_chars = sorted(char_counts.items(), key=lambda x: -x[1])
        top_char = sorted_chars[0]
        if len(sorted_chars) >= 3:
            top3_pct = sum(c[1] for c in sorted_chars[:3]) / sum(c[1] for c in sorted_chars) * 100
            bottom_half = sorted_chars[len(sorted_chars)//2:]
            bottom_pct = sum(c[1] for c in bottom_half) / sum(c[1] for c in sorted_chars) * 100

            if top3_pct > 70:
                critiques.append({
                    "category": "Character Balance",
                    "severity": "warning",
                    "finding": f"Top-heavy cast: the top 3 characters ({', '.join(c[0] for c in sorted_chars[:3])}) "
                               f"account for {top3_pct:.0f}% of all character appearances. "
                               f"The bottom half of the cast shares only {bottom_pct:.0f}%. "
                               f"Consider whether secondary characters have enough development.",
                    "data": {c[0]: c[1] for c in sorted_chars[:10]},
                })
            else:
                critiques.append({
                    "category": "Character Balance",
                    "severity": "strength",
                    "finding": f"Well-distributed screen presence across {len(sorted_chars)} characters. "
                               f"Top 3 at {top3_pct:.0f}% leaves room for the ensemble.",
                    "data": {c[0]: c[1] for c in sorted_chars[:10]},
                })

    # 4. Dialog density
    if num_scenes:
        dialog_per_scene = total_dialog_lines / num_scenes
        silent_pct = scenes_without_dialog / num_scenes * 100

        if silent_pct > 40:
            critiques.append({
                "category": "Dialog Density",
                "severity": "warning",
                "finding": f"{silent_pct:.0f}% of scenes have no dialog at all. "
                           f"Average of {dialog_per_scene:.1f} lines per scene overall. "
                           f"Heavy reliance on visual storytelling or possible indexing gaps.",
                "data": {"avg_lines_per_scene": round(dialog_per_scene, 1), "silent_scenes_pct": round(silent_pct, 1)},
            })
        elif dialog_per_scene > 15:
            critiques.append({
                "category": "Dialog Density",
                "severity": "note",
                "finding": f"Dialog-heavy at {dialog_per_scene:.1f} lines per scene. "
                           f"Only {silent_pct:.0f}% of scenes are purely visual. "
                           f"Consider whether some exposition could be shown rather than told.",
                "data": {"avg_lines_per_scene": round(dialog_per_scene, 1), "silent_scenes_pct": round(silent_pct, 1)},
            })
        else:
            critiques.append({
                "category": "Dialog Density",
                "severity": "strength",
                "finding": f"Balanced dialog density at {dialog_per_scene:.1f} lines per scene. "
                           f"{silent_pct:.0f}% visual-only scenes provide breathing room.",
                "data": {"avg_lines_per_scene": round(dialog_per_scene, 1), "silent_scenes_pct": round(silent_pct, 1)},
            })

    # 5. Visual variety
    if camera_counts:
        dominant_shot = max(camera_counts, key=camera_counts.get)
        shot_pct = camera_counts[dominant_shot] / num_scenes * 100 if num_scenes else 0
        if shot_pct > 50:
            critiques.append({
                "category": "Visual Variety",
                "severity": "warning",
                "finding": f"Repetitive cinematography: '{dominant_shot}' shots used in {shot_pct:.0f}% of scenes. "
                           f"Only {len(camera_counts)} distinct shot types. More variety would enhance visual interest.",
                "data": dict(camera_counts),
            })
        else:
            critiques.append({
                "category": "Visual Variety",
                "severity": "strength",
                "finding": f"Good cinematographic variety with {len(camera_counts)} shot types. "
                           f"No single shot dominates ({dominant_shot} at {shot_pct:.0f}%).",
                "data": dict(camera_counts),
            })

    # 6. Location variety
    if location_counts:
        unique_locations = len(location_counts)
        top_location = max(location_counts, key=location_counts.get)
        top_loc_pct = location_counts[top_location] / num_scenes * 100 if num_scenes else 0
        locations_per_episode = unique_locations / num_episodes if num_episodes else 0

        if locations_per_episode < 3:
            critiques.append({
                "category": "Location Variety",
                "severity": "warning",
                "finding": f"Only {locations_per_episode:.1f} unique locations per episode (avg). "
                           f"'{top_location}' appears in {top_loc_pct:.0f}% of scenes. "
                           f"Limited settings may feel claustrophobic or budget-constrained.",
                "data": {"unique_locations": unique_locations, "top_location": top_location, "top_pct": round(top_loc_pct, 1)},
            })
        else:
            critiques.append({
                "category": "Location Variety",
                "severity": "strength",
                "finding": f"{locations_per_episode:.1f} unique locations per episode across {unique_locations} total. "
                           f"Good spatial variety keeps the world feeling expansive.",
                "data": {"unique_locations": unique_locations, "locations_per_ep": round(locations_per_episode, 1)},
            })

    # 7. Plot momentum
    if plot_sig_counts:
        critical_pct = (plot_sig_counts.get("critical", 0) + plot_sig_counts.get("high", 0)) / num_scenes * 100 if num_scenes else 0
        low_pct = plot_sig_counts.get("low", 0) / num_scenes * 100 if num_scenes else 0

        if low_pct > 50:
            critiques.append({
                "category": "Plot Momentum",
                "severity": "warning",
                "finding": f"{low_pct:.0f}% of scenes rated 'low' plot significance. "
                           f"Only {critical_pct:.0f}% are 'high' or 'critical'. "
                           f"The narrative may drag with too much filler between key moments.",
                "data": dict(plot_sig_counts),
            })
        elif critical_pct > 60:
            critiques.append({
                "category": "Plot Momentum",
                "severity": "note",
                "finding": f"{critical_pct:.0f}% of scenes are 'high' or 'critical' significance. "
                           f"Relentless plot density may exhaust audiences — consider more breathing room.",
                "data": dict(plot_sig_counts),
            })
        else:
            critiques.append({
                "category": "Plot Momentum",
                "severity": "strength",
                "finding": f"Well-balanced plot distribution: {critical_pct:.0f}% high/critical, "
                           f"{low_pct:.0f}% low. Good mix of tension and release.",
                "data": dict(plot_sig_counts),
            })

    # 8. Per-episode consistency
    if len(ep_data) > 1:
        scene_counts = [e["scene_count"] for e in ep_data]
        avg_scenes = sum(scene_counts) / len(scene_counts)
        max_diff = max(abs(c - avg_scenes) for c in scene_counts)
        if max_diff > avg_scenes * 0.5:
            critiques.append({
                "category": "Episode Consistency",
                "severity": "note",
                "finding": f"Significant scene count variance across episodes: "
                           f"range {min(scene_counts)}-{max(scene_counts)} (avg {avg_scenes:.0f}). "
                           f"Inconsistent density may indicate uneven editing or pacing between episodes.",
                "data": {e["label"]: e["scene_count"] for e in ep_data},
            })

    # Summary stats
    strengths = [c for c in critiques if c["severity"] == "strength"]
    warnings = [c for c in critiques if c["severity"] == "warning"]
    notes = [c for c in critiques if c["severity"] == "note"]

    overall_grade = "A" if len(warnings) == 0 else "B+" if len(warnings) <= 1 else "B" if len(warnings) <= 2 else "C+" if len(warnings) <= 3 else "C"

    return {
        "type": "content_critic",
        "title": "Content Critic Report",
        "overall_grade": overall_grade,
        "summary": {
            "total_scenes": num_scenes,
            "total_episodes": num_episodes,
            "total_duration_formatted": f"{int(total_duration // 3600)}h {int((total_duration % 3600) // 60)}m",
            "unique_characters": len(char_counts),
            "unique_locations": len(location_counts),
            "strengths": len(strengths),
            "warnings": len(warnings),
            "notes": len(notes),
        },
        "critiques": critiques,
        "episode_breakdown": ep_data,
    }


MARKETING_EDITOR_PROMPT = """\
You are a senior media marketing strategist and behavioral scientist. You've been given detailed scene-by-scene data from a TV show and must produce a marketing-focused editorial analysis.

Your analysis should be:
- ACTIONABLE: Every observation should connect to a marketing recommendation
- EVIDENCE-BASED: Cite specific neuroscience, psychology, or behavioral economics research where applicable
- HONEST: Praise what works, criticize what doesn't, always with reasoning
- SPECIFIC: Reference actual scenes, characters, and moments from the data

## Show Data Summary

{show_summary}

## Scene-Level Data

{scene_data}

## Required Output

Return a JSON object with this exact structure:
{{
  "overall_assessment": "2-3 sentence executive summary of the show's marketing potential",
  "target_audiences": [
    {{
      "segment": "Audience segment name",
      "appeal_score": 0.8,
      "reasoning": "Why this audience would connect",
      "hook": "The specific marketing hook for this segment"
    }}
  ],
  "observations": [
    {{
      "category": "Category name (e.g., Emotional Hooks, Character Relatability, Binge Architecture, Social Clip-ability, Villain Effectiveness, Stakes & Tension, Opening Hook, Cliffhanger Quality)",
      "type": "strength|opportunity|concern",
      "observation": "What you observed in the data",
      "marketing_implication": "What this means for marketing and audience engagement",
      "evidence": "Neuroscience, psychology, or behavioral economics citation supporting this point (be specific: author, study, or principle name)",
      "recommendation": "Specific actionable recommendation",
      "scenes_referenced": ["Scene descriptions or timestamps that support this"]
    }}
  ],
  "clip_recommendations": [
    {{
      "purpose": "Social media teaser|Trailer moment|Character intro|Emotional hook|Action highlight",
      "description": "What the clip should contain",
      "timestamp_hint": "Approximate scene/timestamp from the data",
      "platform": "TikTok|Instagram Reels|YouTube|Twitter/X|All",
      "psychological_hook": "Why this clip works psychologically"
    }}
  ],
  "binge_analysis": {{
    "binge_score": 0.8,
    "cliffhanger_effectiveness": "Assessment of episode endings",
    "pacing_for_retention": "Whether pacing encourages continued viewing",
    "drop_off_risks": ["Points where viewers might stop watching and why"]
  }},
  "competitive_positioning": "How this show positions against similar content in the market"
}}

Be thorough but concise. Every observation must be grounded in the actual scene data provided. For neuroscience/psychology citations, reference real research principles (e.g., Zeigarnik effect for cliffhangers, mere exposure effect for character familiarity, peak-end rule for episode structure, narrative transportation theory, parasocial relationships, dopamine reward scheduling)."""


async def build_marketing_editor(episodes_scenes: dict) -> dict:
    """AI-powered marketing editorial analysis using Gemini.

    Sends aggregated scene data to Gemini with a specialized marketing
    analysis prompt. Returns structured observations with evidence-based
    recommendations and neuroscience citations.
    """
    # Build show summary
    all_chars = defaultdict(int)
    all_tones = defaultdict(int)
    all_locations = set()
    total_scenes = 0
    total_duration = 0
    ep_summaries = []

    for ep_id, (ep, scenes) in episodes_scenes.items():
        total_scenes += len(scenes)
        ep_dur = sum(s.duration for s in scenes)
        total_duration += ep_dur

        ep_chars = set()
        ep_tones = defaultdict(int)
        for scene in scenes:
            for c in (scene.characters_present or []):
                name = c.get("name", "Unknown")
                all_chars[name] += 1
                ep_chars.add(name)
            if scene.tone:
                all_tones[scene.tone] += 1
                ep_tones[scene.tone] += 1
            if scene.location:
                all_locations.add(scene.location)

        ep_summaries.append(
            f"{_ep_label(ep)} '{ep.title}': {len(scenes)} scenes, "
            f"{int(ep_dur//60)}min, chars: {', '.join(sorted(ep_chars)[:8])}, "
            f"tones: {dict(ep_tones)}"
        )

    show_summary = (
        f"Episodes: {len(episodes_scenes)}\n"
        f"Total scenes: {total_scenes}\n"
        f"Total duration: {int(total_duration//3600)}h {int((total_duration%3600)//60)}m\n"
        f"Characters ({len(all_chars)}): {', '.join(f'{n} ({c} scenes)' for n, c in sorted(all_chars.items(), key=lambda x: -x[1])[:15])}\n"
        f"Tones: {dict(all_tones)}\n"
        f"Locations: {len(all_locations)} unique\n"
        f"Episodes:\n" + "\n".join(f"  {s}" for s in ep_summaries)
    )

    # Build scene data — send first 3 and last 2 scenes per episode (key moments)
    # plus any high/critical significance scenes
    scene_entries = []
    for ep_id, (ep, scenes) in episodes_scenes.items():
        key_scenes = []
        for i, scene in enumerate(scenes):
            is_key = (
                i < 3 or i >= len(scenes) - 2  # Opening and closing scenes
                or scene.plot_significance in ("high", "critical")
            )
            if is_key and len(key_scenes) < 15:  # Cap per episode
                chars = [c.get("name", "?") for c in (scene.characters_present or [])]
                dialog_sample = ""
                for d in (scene.key_dialog or [])[:3]:
                    dialog_sample += f'  {d.get("speaker","?")}: "{d.get("quote","")[:80]}"\n'

                key_scenes.append(
                    f"  [{_fmt_time(scene.start_timestamp)}-{_fmt_time(scene.end_timestamp)}] "
                    f"Tone:{scene.tone} Pacing:{scene.scene_pacing} Plot:{scene.plot_significance}\n"
                    f"    Location: {scene.location or '?'}\n"
                    f"    Characters: {', '.join(chars)}\n"
                    f"    Description: {(scene.description_text or '')[:200]}\n"
                    f"    Mood: {scene.mood_ambience or '?'}\n"
                    f"    Emotional arc: {scene.emotional_arc or '?'}\n"
                    + (f"    Dialog:\n{dialog_sample}" if dialog_sample else "")
                    + (f"    Motivations: {scene.character_motivations_feelings[:150]}\n" if scene.character_motivations_feelings else "")
                )

        scene_entries.append(f"\n{_ep_label(ep)} — {ep.title}:\n" + "\n".join(key_scenes))

    scene_data = "\n".join(scene_entries)

    # Call Gemini
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return {
            "type": "marketing_editor",
            "title": "Marketing Editor Analysis",
            "error": "GEMINI_API_KEY not set — cannot run AI analysis",
        }

    try:
        from google import genai
        from google.genai import types

        client = genai.Client(api_key=api_key)

        prompt = MARKETING_EDITOR_PROMPT.format(
            show_summary=show_summary,
            scene_data=scene_data[:15000],  # Cap to avoid token limits
        )

        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.4,
                response_mime_type="application/json",
            ),
        )

        result = json.loads(response.text)

        return {
            "type": "marketing_editor",
            "title": "Marketing Editor Analysis",
            "analysis": result,
        }

    except Exception as e:
        logger.error(f"Marketing Editor Gemini call failed: {e}")
        return {
            "type": "marketing_editor",
            "title": "Marketing Editor Analysis",
            "error": f"AI analysis failed: {str(e)[:200]}",
        }


BUILDERS = {
    "plot_summary": build_plot_summary,
    "character_summaries": build_character_summaries,
    "character_motivations": build_character_motivations,
    "two_column_script": build_two_column_script,
    "scene_breakdown": build_scene_breakdown,
    "dialog_only": build_dialog_only,
    "visual_descriptions": build_visual_descriptions,
    "content_critic": build_content_critic,
    "marketing_editor": build_marketing_editor,
}


# ─── DOCX Generation ───────────────────────────────────────

def _add_cell_borders(cell, color="CCCCCC"):
    """Add thin borders to a table cell."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = tcPr.makeelement(qn('w:tcBorders'), {})
        tcPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:space'): '0', qn('w:color'): color,
        })
        borders.append(el)


def _build_docx(show: Show, sections: list[dict], scope_label: str) -> io.BytesIO:
    """Assemble all report sections into a single DOCX."""
    from docx.shared import Cm
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    import datetime

    doc = DocxDocument()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Style headings
    for level, (size, color) in enumerate([(24, 0x2E75B6), (18, 0x333333), (14, 0x555555)], 1):
        hstyle = doc.styles[f'Heading {level}']
        hstyle.font.name = 'Arial'
        hstyle.font.size = Pt(size)
        hstyle.font.color.rgb = RGBColor((color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF)
        hstyle.font.bold = True
        if level == 1:
            hstyle.paragraph_format.space_before = Pt(24)
            hstyle.paragraph_format.space_after = Pt(12)

    # ── Title Page ──
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(show.name.upper())
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("━" * 30)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(scope_label)
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    for _ in range(3):
        doc.add_paragraph()

    # Table of contents
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONTENTS")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r.bold = True

    doc.add_paragraph()
    for i, section in enumerate(sections, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{i}. {section.get('title', 'Section')}")
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Generated by Narralytica — {datetime.date.today().isoformat()}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    r.italic = True

    doc.add_page_break()

    for section in sections:
        stype = section.get("type", "")

        if stype == "plot_summary":
            doc.add_heading("Plot Summary", level=1)
            for ep in section.get("episodes", []):
                doc.add_heading(f"{ep['episode_label']} — {ep['episode_title']}", level=2)
                if ep.get("narrative"):
                    for para_text in ep["narrative"].split("\n\n"):
                        if para_text.strip():
                            p = doc.add_paragraph(para_text.strip())
                            p.runs[0].font.size = Pt(11)

                # Scene summaries as appendix
                if ep.get("scene_summaries"):
                    doc.add_heading("Scene-by-Scene", level=3)
                    for i, summary in enumerate(ep["scene_summaries"], 1):
                        p = doc.add_paragraph()
                        r = p.add_run(f"Scene {i}: ")
                        r.bold = True
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                        r = p.add_run(str(summary))
                        r.font.size = Pt(9)
            doc.add_page_break()

        elif stype == "character_summaries":
            doc.add_heading("Character Summaries", level=1)
            for char in section.get("characters", []):
                doc.add_heading(char["name"], level=2)
                p = doc.add_paragraph()
                p.add_run(f"Scenes: {char['scene_count']}  |  "
                          f"Duration: {char['duration_formatted']}  |  "
                          f"Speaking lines: {char['speaking_lines']}  |  "
                          f"Dominant tone: {char.get('dominant_tone', 'N/A')}")
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                if char.get("description"):
                    p = doc.add_paragraph(char["description"])
                    p.runs[0].font.size = Pt(11)

                if char.get("episodes"):
                    p = doc.add_paragraph(f"Episodes: {', '.join(char['episodes'])}")
                    p.runs[0].font.size = Pt(9)

                if char.get("key_quotes"):
                    for q in char["key_quotes"]:
                        p = doc.add_paragraph()
                        r = p.add_run(f'"{q}"')
                        r.italic = True
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_page_break()

        elif stype == "character_motivations":
            doc.add_heading("Character Motivations", level=1)
            p = doc.add_paragraph("What each character wants and feels, scene by scene.")
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            doc.add_paragraph()

            for name, entries in section.get("characters", {}).items():
                doc.add_heading(f"{name} ({len(entries)} scenes)", level=2)

                for entry in entries:
                    # Scene context line
                    p = doc.add_paragraph()
                    r = p.add_run(f"Scene {entry.get('scene', '?')}")
                    r.bold = True
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                    r = p.add_run(f"  {entry['time']}  |  {entry.get('location', '')}  |  {entry.get('tone', '')}")
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                    # Motivation text
                    motivation = entry.get("motivation", entry.get("text", ""))
                    if motivation:
                        p = doc.add_paragraph()
                        r = p.add_run(motivation)
                        r.font.size = Pt(10)
                        p.paragraph_format.left_indent = Pt(18)
                    doc.add_paragraph()  # spacer

                doc.add_paragraph()  # spacer between characters
            doc.add_page_break()

        elif stype == "two_column_script":
            doc.add_heading("Two-Column Script", level=1)
            for ep in section.get("episodes", []):
                doc.add_heading(f"{ep['episode_label']} — {ep['episode_title']}", level=2)

                for row in ep.get("rows", []):
                    # Scene header
                    p = doc.add_paragraph()
                    r = p.add_run(f"Scene {row['scene_number']} — {row['time']}")
                    r.bold = True
                    r.font.size = Pt(10)
                    if row.get("location"):
                        r = p.add_run(f"  |  {row['location']}")
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                    # Two-column table
                    table = doc.add_table(rows=1, cols=2)
                    table.autofit = False
                    table.columns[0].width = Inches(3.25)
                    table.columns[1].width = Inches(3.25)

                    # Header row
                    hdr = table.rows[0]
                    hdr.cells[0].text = "DIALOG / AUDIO"
                    hdr.cells[1].text = "VISUALS / ACTION"
                    for cell in hdr.cells:
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                        _add_cell_borders(cell, "888888")

                    # Content row
                    content_row = table.add_row()
                    left_cell = content_row.cells[0]
                    right_cell = content_row.cells[1]
                    _add_cell_borders(left_cell)
                    _add_cell_borders(right_cell)

                    # LEFT: Dialog
                    left_cell.text = ""
                    for line in row["left"]["dialog"]:
                        p = left_cell.add_paragraph()
                        r = p.add_run(f"{line['speaker']}: ")
                        r.bold = True
                        r.font.size = Pt(9)
                        r = p.add_run(line["text"])
                        r.font.size = Pt(9)
                        p.paragraph_format.space_after = Pt(2)

                    for audio_line in row["left"]["audio"]:
                        p = left_cell.add_paragraph()
                        r = p.add_run(audio_line)
                        r.italic = True
                        r.font.size = Pt(8)
                        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

                    # RIGHT: Visuals
                    right_cell.text = ""
                    right = row["right"]
                    if right.get("description"):
                        p = right_cell.add_paragraph()
                        r = p.add_run(right["description"])
                        r.font.size = Pt(9)
                        p.paragraph_format.space_after = Pt(4)

                    if right.get("actions"):
                        p = right_cell.add_paragraph()
                        r = p.add_run(right["actions"])
                        r.italic = True
                        r.font.size = Pt(9)
                        p.paragraph_format.space_after = Pt(4)

                    meta = []
                    if right.get("camera"):
                        meta.append(f"Camera: {right['camera']}")
                    if right.get("lighting"):
                        meta.append(f"Lighting: {right['lighting']}")
                    if right.get("composition"):
                        meta.append(f"Comp: {right['composition']}")
                    if meta:
                        p = right_cell.add_paragraph()
                        r = p.add_run(" | ".join(meta))
                        r.font.size = Pt(8)
                        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                    doc.add_paragraph()  # Spacer

            doc.add_page_break()

        elif stype == "scene_breakdown":
            doc.add_heading("Scene-by-Scene Breakdown", level=1)
            for ep in section.get("episodes", []):
                doc.add_heading(f"{ep['episode_label']} — {ep['episode_title']}", level=2)
                for scene in ep.get("scenes", []):
                    doc.add_heading(
                        f"Scene {scene['scene_number']} — {scene['time']} ({scene['duration']})",
                        level=3,
                    )
                    if scene.get("characters"):
                        doc.add_paragraph(f"Characters: {', '.join(scene['characters'])}")
                    meta = []
                    if scene.get("tone"):
                        meta.append(f"Tone: {scene['tone']}")
                    if scene.get("mood"):
                        meta.append(f"Mood: {scene['mood']}")
                    if scene.get("pacing"):
                        meta.append(f"Pacing: {scene['pacing']}")
                    if scene.get("plot_significance"):
                        meta.append(f"Plot: {scene['plot_significance']}")
                    if meta:
                        p = doc.add_paragraph(" | ".join(meta))
                        p.runs[0].font.size = Pt(9)
                        p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                    if scene.get("description"):
                        doc.add_paragraph(scene["description"])
                    if scene.get("actions"):
                        p = doc.add_paragraph(scene["actions"])
                        p.runs[0].italic = True
            doc.add_page_break()

        elif stype == "dialog_only":
            doc.add_heading("Dialog Transcript", level=1)
            for ep in section.get("episodes", []):
                doc.add_heading(f"{ep['episode_label']} — {ep['episode_title']}", level=2)
                for scene in ep.get("scenes", []):
                    p = doc.add_paragraph()
                    r = p.add_run(f"[{scene['scene_time']}] {scene['location']}")
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

                    for line in scene["dialog"]:
                        p = doc.add_paragraph()
                        r = p.add_run(f"{line['speaker']}: ")
                        r.bold = True
                        r.font.size = Pt(10)
                        r = p.add_run(line["text"])
                        r.font.size = Pt(10)
                        p.paragraph_format.space_after = Pt(1)
            doc.add_page_break()

        elif stype == "visual_descriptions":
            doc.add_heading("Visual Descriptions", level=1)
            for ep in section.get("episodes", []):
                doc.add_heading(f"{ep['episode_label']} — {ep['episode_title']}", level=2)
                for scene in ep.get("scenes", []):
                    doc.add_heading(
                        f"Scene {scene['scene_number']} — {scene['time']}",
                        level=3,
                    )
                    if scene.get("location"):
                        p = doc.add_paragraph(f"Location: {scene['location']}")
                        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                    if scene.get("description"):
                        doc.add_paragraph(scene["description"])
                    if scene.get("actions"):
                        p = doc.add_paragraph(scene["actions"])
                        p.runs[0].italic = True
                    meta = []
                    if scene.get("camera"):
                        meta.append(f"Camera: {scene['camera']}")
                    if scene.get("lighting"):
                        meta.append(f"Lighting: {scene['lighting']}")
                    if scene.get("visual_style"):
                        meta.append(f"Style: {scene['visual_style']}")
                    if meta:
                        p = doc.add_paragraph(" | ".join(meta))
                        p.runs[0].font.size = Pt(9)
                        p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        elif stype == "content_critic":
            SEVERITY_COLORS = {
                "strength": RGBColor(0x27, 0xAE, 0x60),
                "warning": RGBColor(0xE7, 0x4C, 0x3C),
                "note": RGBColor(0xF3, 0x9C, 0x12),
            }
            SEVERITY_ICONS = {"strength": "+", "warning": "!", "note": "~"}

            doc.add_heading("Content Critic Report", level=1)

            # Overall grade
            grade = section.get("overall_grade", "?")
            summary = section.get("summary", {})
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"Overall Grade: {grade}")
            r.bold = True
            r.font.size = Pt(36)
            r.font.color.rgb = SEVERITY_COLORS.get("strength") if grade.startswith("A") else SEVERITY_COLORS.get("warning") if grade.startswith("C") else RGBColor(0x2E, 0x75, 0xB6)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(
                f"{summary.get('total_episodes', '?')} episodes | "
                f"{summary.get('total_scenes', '?')} scenes | "
                f"{summary.get('total_duration_formatted', '?')} | "
                f"{summary.get('unique_characters', '?')} characters | "
                f"{summary.get('strengths', 0)} strengths, {summary.get('warnings', 0)} warnings"
            )
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            doc.add_paragraph()

            # Critiques
            for critique in section.get("critiques", []):
                severity = critique.get("severity", "note")
                icon = SEVERITY_ICONS.get(severity, "?")
                color = SEVERITY_COLORS.get(severity, RGBColor(0x88, 0x88, 0x88))

                p = doc.add_paragraph()
                r = p.add_run(f"[{icon}] {critique['category']}")
                r.bold = True
                r.font.size = Pt(12)
                r.font.color.rgb = color

                p = doc.add_paragraph(critique["finding"])
                p.runs[0].font.size = Pt(10)

                if critique.get("data"):
                    p = doc.add_paragraph()
                    r = p.add_run(f"Data: {', '.join(f'{k}: {v}' for k, v in critique['data'].items())}")
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
                    r.italic = True

                doc.add_paragraph()

            # Episode breakdown table
            if section.get("episode_breakdown"):
                doc.add_heading("Per-Episode Breakdown", level=2)
                eps = section["episode_breakdown"]
                table = doc.add_table(rows=1, cols=6)
                table.autofit = True
                for i, hdr_text in enumerate(["Episode", "Scenes", "Characters", "Dialog Lines", "Tone", "Pacing Shifts"]):
                    table.rows[0].cells[i].text = hdr_text
                    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                    table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(8)

                for ep in eps:
                    row = table.add_row()
                    row.cells[0].text = f"{ep['label']} {ep['title']}"
                    row.cells[1].text = str(ep["scene_count"])
                    row.cells[2].text = str(ep["unique_chars"])
                    row.cells[3].text = str(ep["dialog_lines"])
                    row.cells[4].text = ep["dominant_tone"]
                    row.cells[5].text = str(ep["pacing_shifts"])
                    for cell in row.cells:
                        cell.paragraphs[0].runs[0].font.size = Pt(8)

            doc.add_page_break()

        elif stype == "marketing_editor":
            analysis = section.get("analysis", {})
            if section.get("error"):
                doc.add_heading("Marketing Editor Analysis", level=1)
                doc.add_paragraph(f"Error: {section['error']}")
            else:
                doc.add_heading("Marketing Editor Analysis", level=1)

                # Overall assessment
                if analysis.get("overall_assessment"):
                    p = doc.add_paragraph()
                    r = p.add_run(analysis["overall_assessment"])
                    r.font.size = Pt(12)
                    r.italic = True
                    doc.add_paragraph()

                # Target audiences
                if analysis.get("target_audiences"):
                    doc.add_heading("Target Audiences", level=2)
                    for aud in analysis["target_audiences"]:
                        p = doc.add_paragraph()
                        r = p.add_run(f"{aud.get('segment', '?')} ")
                        r.bold = True
                        r.font.size = Pt(11)
                        score = aud.get("appeal_score", 0)
                        r = p.add_run(f"(Appeal: {int(score * 100)}%)")
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60) if score >= 0.7 else RGBColor(0xF3, 0x9C, 0x12)
                        doc.add_paragraph(aud.get("reasoning", ""))
                        if aud.get("hook"):
                            p = doc.add_paragraph()
                            r = p.add_run(f"Hook: {aud['hook']}")
                            r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                            r.font.size = Pt(10)
                        doc.add_paragraph()

                # Observations
                if analysis.get("observations"):
                    doc.add_heading("Observations & Recommendations", level=2)
                    TYPE_COLORS = {
                        "strength": RGBColor(0x27, 0xAE, 0x60),
                        "opportunity": RGBColor(0x2E, 0x75, 0xB6),
                        "concern": RGBColor(0xE7, 0x4C, 0x3C),
                    }
                    for obs in analysis["observations"]:
                        p = doc.add_paragraph()
                        obs_type = obs.get("type", "note")
                        r = p.add_run(f"[{obs_type.upper()}] {obs.get('category', '?')}")
                        r.bold = True
                        r.font.size = Pt(11)
                        r.font.color.rgb = TYPE_COLORS.get(obs_type, RGBColor(0x88, 0x88, 0x88))

                        doc.add_paragraph(obs.get("observation", ""))

                        if obs.get("marketing_implication"):
                            p = doc.add_paragraph()
                            r = p.add_run("Marketing implication: ")
                            r.bold = True
                            r.font.size = Pt(10)
                            r = p.add_run(obs["marketing_implication"])
                            r.font.size = Pt(10)

                        if obs.get("evidence"):
                            p = doc.add_paragraph()
                            r = p.add_run(f"Evidence: {obs['evidence']}")
                            r.italic = True
                            r.font.size = Pt(9)
                            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                        if obs.get("recommendation"):
                            p = doc.add_paragraph()
                            r = p.add_run(f"Recommendation: {obs['recommendation']}")
                            r.font.size = Pt(10)
                            r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

                        doc.add_paragraph()

                # Clip recommendations
                if analysis.get("clip_recommendations"):
                    doc.add_heading("Social Media Clip Recommendations", level=2)
                    for clip in analysis["clip_recommendations"]:
                        p = doc.add_paragraph()
                        r = p.add_run(f"{clip.get('purpose', '?')}")
                        r.bold = True
                        r = p.add_run(f" — {clip.get('platform', '?')}")
                        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                        doc.add_paragraph(clip.get("description", ""))
                        if clip.get("psychological_hook"):
                            p = doc.add_paragraph()
                            r = p.add_run(f"Why it works: {clip['psychological_hook']}")
                            r.italic = True
                            r.font.size = Pt(9)
                        doc.add_paragraph()

                # Binge analysis
                binge = analysis.get("binge_analysis", {})
                if binge:
                    doc.add_heading("Binge-Watching Analysis", level=2)
                    score = binge.get("binge_score", 0)
                    p = doc.add_paragraph()
                    r = p.add_run(f"Binge Score: {int(score * 100)}%")
                    r.bold = True
                    r.font.size = Pt(14)
                    if binge.get("cliffhanger_effectiveness"):
                        doc.add_paragraph(f"Cliffhangers: {binge['cliffhanger_effectiveness']}")
                    if binge.get("pacing_for_retention"):
                        doc.add_paragraph(f"Pacing: {binge['pacing_for_retention']}")
                    if binge.get("drop_off_risks"):
                        doc.add_paragraph("Drop-off risks:")
                        for risk in binge["drop_off_risks"]:
                            doc.add_paragraph(risk, style='List Bullet')

                # Competitive positioning
                if analysis.get("competitive_positioning"):
                    doc.add_heading("Competitive Positioning", level=2)
                    doc.add_paragraph(analysis["competitive_positioning"])

            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── Main Endpoint ──────────────────────────────────────────

@router.post("/generate")
async def generate_report(body: ReportRequest, db: AsyncSession = Depends(get_db)):
    """Generate a report from indexed scene data.

    Returns JSON preview or DOCX download depending on format param.
    """
    # Validate
    invalid = set(body.report_types) - VALID_REPORT_TYPES
    if invalid:
        raise HTTPException(400, f"Invalid report types: {invalid}")

    if not body.report_types:
        raise HTTPException(400, "Select at least one report type")

    # Fetch data
    if body.scope == "episode":
        if not body.episode_id:
            raise HTTPException(400, "episode_id required for episode scope")
        show, episodes_scenes = await _get_episode_data(body.episode_id, db)
        ep = list(episodes_scenes.values())[0][0]
        scope_label = f"{_ep_label(ep)} — {ep.title}"
    elif body.scope == "season":
        if not body.show_id or body.season is None:
            raise HTTPException(400, "show_id and season required for season scope")
        show, episodes_scenes = await _get_season_data(body.show_id, body.season, db)
        scope_label = f"Season {body.season} ({len(episodes_scenes)} episodes)"
    else:
        raise HTTPException(400, f"Invalid scope: {body.scope}")

    # Build sections
    sections = []
    for rt in body.report_types:
        builder = BUILDERS.get(rt)
        if builder:
            import asyncio
            result = builder(episodes_scenes)
            # Handle async builders (marketing_editor uses Gemini)
            if asyncio.iscoroutine(result):
                result = await result
            sections.append(result)

    if body.format == "docx":
        buf = _build_docx(show, sections, scope_label)
        filename = _safe_filename(show.name, scope_label, "docx")
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # Preview
    return {
        "title": f"{show.name} — {scope_label}",
        "sections": sections,
    }
